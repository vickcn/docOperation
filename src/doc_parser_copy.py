from docx import Document
import json
import os
from typing import Dict, Any
import logging
from pathlib import Path
from docx.shared import Pt, Twips, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

def merge_chord_runs(handler, runs):
    """合并和弦相关的 runs，例如将 C 和 # 合并为 C#，而数字作为单独的修饰符"""
    if not runs:
        return runs
        
    merged_runs = []
    i = 0
    while i < len(runs):
        current_run = runs[i].copy()  # 复制当前 run
        
        # 检查是否是可能的和弦开始（大写字母）
        if (current_run.get('text', '').strip() and 
            current_run['text'][0].isupper() and 
            len(current_run['text']) == 1):
            
            # 只检查下一个 run 是否是升降号
            next_idx = i + 1
            # print('i', i)
            if next_idx < len(runs):
                next_run = runs[next_idx]
                next_text = next_run.get('text', '').strip()
                
                # 如果是升降号，合并到当前 run
                # if next_text in ['#', 'b']:
                print('next_text', next_text)
                if next_text.find('#')>-1 or next_text.find('b')>-1:
                    # 合并文本
                    current_run['text'] += next_text[0]
                    # 记录升降号的格式信息
                    if not 'accidentals' in current_run:
                        current_run['accidentals'] = []
                    current_run['accidentals'].append({
                        'text': next_text,
                        'superscript': next_run.get('superscript', False),
                        'font_size': next_run.get('font_size')
                    })
                    print('current_run', current_run['text'])
                    # i = next_idx + 1  # 跳过已合并的升降号
                # else:
                    # i += 1
            #     i += 1
            # else:
            #     i += 1
            i += 1
        elif(current_run.get('text', '').strip() and 
             current_run['text'][0] in ['#','b']):
            current_run['text'] = current_run['text'][1:]
            i += 1
        else:
            i += 1
            
        merged_runs.append(current_run)
        
    return merged_runs 

class DocParser:
    def __init__(self, logger=None):
        self.logger = logger or logging.getLogger(__name__)
        self._current_paragraph_runs = []  # 用于临时存储当前段落的所有 runs
        
    def _should_merge_with_next(self, current_run, next_run) -> bool:
        """判断当前 run 是否应该与下一个 run 合并"""
        current_text = current_run.get('text', '').strip()
        next_text = next_run.get('text', '').strip()
        
        # 检查是否是和弦情况（大写字母后跟升降号）
        is_chord_start = (current_text and 
                         len(current_text) == 1 and 
                         current_text[0].isupper())
        is_accidental = next_text in ['#', 'b']
        
        print(f"检查是否需要合并: current='{current_text}', next='{next_text}'")
        print(f"is_chord_start={is_chord_start}, is_accidental={is_accidental}")
        
        return is_chord_start and is_accidental

    def _extract_run_formatting(self, run) -> Dict[str, Any]:
        """提取文本运行的格式信息"""
        run_format = {
            "text": run.text,
            "bold": run.bold,
            "italic": run.italic,
            "underline": run.underline,
            "font_size": self._convert_size_to_pt(run.font.size),
            "font_name": run.font.name,
            "style": run.style.name if run.style else None,
            "hyperlink": self._extract_hyperlink(run),
            "subscript": run._element.rPr.vertAlign is not None and run._element.rPr.vertAlign.val == 'subscript',
            "superscript": run._element.rPr.vertAlign is not None and run._element.rPr.vertAlign.val == 'superscript'
        }

        # 提取字体详细信息
        rPr = run._element.rPr
        if rPr is not None:
            rFonts = rPr.rFonts
            if rFonts is not None:
                # 提取各种字体设置
                run_format["font_ascii"] = rFonts.get(qn('w:ascii'))
                run_format["font_east_asia"] = rFonts.get(qn('w:eastAsia'))
                run_format["font_h_ansi"] = rFonts.get(qn('w:hAnsi'))
                run_format["font_cs"] = rFonts.get(qn('w:cs'))
                run_format["font_hint"] = rFonts.get(qn('w:hint'))

        # 提取颜色信息
        if run.font.color.rgb:
            run_format["color"] = run.font.color.rgb
        elif run._element.rPr.color is not None:
            run_format["color"] = run._element.rPr.color.val

        # 提取突出显示颜色
        if run._element.rPr.highlight is not None:
            run_format["highlight_color"] = run._element.rPr.highlight.val

        # 存储当前 run 的格式信息
        self._current_paragraph_runs.append(run_format)
        
        # 检查是否需要与前一个 run 合并
        if len(self._current_paragraph_runs) >= 2:
            prev_run = self._current_paragraph_runs[-2]
            current_run = self._current_paragraph_runs[-1]
            
            # 检查是否是和弦情况（大写字母后跟升降号）
            prev_text = prev_run.get('text', '').strip()
            current_text = current_run.get('text', '').strip()
            
            # 检查是否是大写字母后跟升降号
            is_chord_start = (prev_text and 
                            len(prev_text) == 1 and 
                            prev_text[0].isupper())
            is_accidental = current_text in ['#', 'b']
            
            if is_chord_start and is_accidental:
                # 合并文本
                prev_run['text'] += current_text
                # 记录升降号的格式信息
                if not 'accidentals' in prev_run:
                    prev_run['accidentals'] = []
                prev_run['accidentals'].append({
                    'text': current_text,
                    'superscript': current_run.get('superscript', False),
                    'font_size': current_run.get('font_size'),
                    'font_ascii': current_run.get('font_ascii'),
                    'font_east_asia': current_run.get('font_east_asia'),
                    'font_h_ansi': current_run.get('font_h_ansi'),
                    'font_cs': current_run.get('font_cs'),
                    'font_hint': current_run.get('font_hint')
                })
                # 移除当前 run（因为已经合并到前一个 run 中）
                self._current_paragraph_runs.pop()
                print(f"合并和弦: {prev_run['text']}, 升降号格式: {prev_run.get('accidentals')}")
                return None  # 返回 None 表示这个 run 已经被合并

        return run_format

    def parse_paragraph(self, paragraph):
        """解析段落，提取格式信息"""
        para_format = {}
        
        # 提取段落格式
        if paragraph.paragraph_format.line_spacing:
            para_format['line_spacing'] = paragraph.paragraph_format.line_spacing
            
        if paragraph.paragraph_format.space_before:
            para_format['space_before'] = paragraph.paragraph_format.space_before
            
        if paragraph.paragraph_format.space_after:
            para_format['space_after'] = paragraph.paragraph_format.space_after
            
        if paragraph.paragraph_format.first_line_indent:
            para_format['first_line_indent'] = paragraph.paragraph_format.first_line_indent
            
        # 重置当前段落的 runs 列表
        self._current_paragraph_runs = []
        
        # 提取 runs 信息
        runs = []
        for run in paragraph.runs:
            run_info = self._extract_run_formatting(run)
            if run_info is not None:  # 只添加未被合并的 run
                runs.append(run_info)
        # 合并和弦相关的 runs
        runs = self.merge_chord_runs(runs)
        return {
            'text': paragraph.text,
            'style': paragraph.style.name,
            'format': para_format,
            'runs': runs
        }

    def parse_docx(self, file_path: str) -> Dict[str, Any]:
        """
        解析docx文件並返回結構化數據
        """
        try:
            doc = Document(file_path)
            result = {
                "metadata": self._extract_metadata(doc),
                "sections": self._extract_section_properties(doc),
                "paragraphs": self._extract_paragraphs(doc),
                "tables": self._extract_tables(doc),
                "images": self._extract_images(doc),
                "styles": self._extract_styles(doc)
            }
            self.logger.info(f"Successfully parsed document: {file_path}")
            return result
        except Exception as e:
            self.logger.error(f"Error parsing document: {str(e)}")
            raise

    def _extract_styles(self, doc: Document) -> Dict[str, Any]:
        """提取文档样式信息"""
        styles = {}
        for style in doc.styles:
            if style.type == 1:  # 段落样式
                styles[style.name] = {
                    "name": style.name,
                    "font_name": style.font.name if style.font else None,
                    "font_size": self._convert_size_to_pt(style.font.size) if style.font else None,
                    "bold": style.font.bold if style.font else None,
                    "italic": style.font.italic if style.font else None
                }
        return styles

    def _extract_metadata(self, doc: Document) -> Dict[str, Any]:
        """提取文档元数据"""
        return {
            "core_properties": {
                "title": doc.core_properties.title,
                "author": doc.core_properties.author,
                "created": str(doc.core_properties.created),
                "modified": str(doc.core_properties.modified)
            }
        }

    def _convert_size_to_pt(self, size) -> float:
        """将字体大小转换为磅值"""
        if isinstance(size, Pt):
            return float(size.pt)
        elif isinstance(size, Twips):
            return float(size.pt)
        elif size is not None:
            try:
                return float(size) / 20  # 转换默认单位到磅值
            except (ValueError, TypeError):
                return 12.0  # 默认12磅
        return 12.0  # 默认12磅

    def _extract_hyperlink(self, run) -> Dict[str, Any]:
        """提取超連結信息"""
        # 檢查父元素是否是超連結
        parent = run._element.getparent()
        if parent.tag.endswith('hyperlink'):
            # 獲取文檔部分
            part = run._parent._parent.part
            # 獲取關係ID
            rel_id = parent.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
            if rel_id and rel_id in part.rels:
                rel = part.rels[rel_id]
                return {
                    "url": rel.target_ref,
                    "text": run.text
                }
        return None

    def _extract_paragraph_format(self, paragraph) -> Dict[str, Any]:
        """提取段落格式信息"""
        pf = paragraph.paragraph_format
        format_data = {
            "alignment": paragraph.alignment,
            "line_spacing": pf.line_spacing,
            "space_before": pf.space_before.pt if pf.space_before else None,
            "space_after": pf.space_after.pt if pf.space_after else None,
            "first_line_indent": pf.first_line_indent.pt if pf.first_line_indent else None,
            "left_indent": pf.left_indent.pt if pf.left_indent else None,
            "right_indent": pf.right_indent.pt if pf.right_indent else None,
            "keep_together": pf.keep_together,
            "keep_with_next": pf.keep_with_next,
            "page_break_before": pf.page_break_before,
            "widow_control": pf.widow_control
        }
        return {k: v for k, v in format_data.items() if v is not None}

    def _extract_paragraphs(self, doc: Document) -> list:
        """提取段落内容"""
        paragraphs = []
        for para in doc.paragraphs:
            para_info = self.parse_paragraph(para)
            paragraphs.append(para_info)
        return paragraphs

    def _extract_tables(self, doc: Document) -> list:
        """提取表格内容"""
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    # 确保正确处理中文字符
                    text = cell.text
                    if not isinstance(text, str):
                        text = str(text, 'utf-8', errors='replace')
                    row_data.append(text)
                table_data.append(row_data)
            tables.append(table_data)
        return tables

    def _extract_images(self, doc: Document) -> list:
        """提取图片信息"""
        images = []
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                images.append({
                    "id": rel.rId,
                    "type": rel.target_ref.split(".")[-1]
                })
        return images

    def _extract_section_properties(self, doc: Document) -> Dict[str, Any]:
        """提取文檔節的屬性，包括分欄設置"""
        sections = []
        for section in doc.sections:
            # 獲取分欄信息
            cols_element = section._sectPr.xpath("./w:cols")
            cols_info = {
                "count": "1",
                "space": "0"
            }
            
            if cols_element:
                cols = cols_element[0]
                cols_info["count"] = cols.get(qn("w:num")) or "1"
                cols_info["space"] = cols.get(qn("w:space")) or "0"
            
            section_props = {
                "start_type": section.start_type,
                "page_height": float(section.page_height.pt),
                "page_width": float(section.page_width.pt),
                "left_margin": float(section.left_margin.pt),
                "right_margin": float(section.right_margin.pt),
                "top_margin": float(section.top_margin.pt),
                "bottom_margin": float(section.bottom_margin.pt),
                "header_distance": float(section.header_distance.pt),
                "footer_distance": float(section.footer_distance.pt),
                "orientation": section.orientation,
                "columns": cols_info
            }
            sections.append(section_props)
        return sections

    def save_to_json(self, data: Dict[str, Any], output_path: str) -> None:
        """将解析结果保存为JSON文件"""
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=4)
            self.logger.info(f"Successfully saved JSON to: {output_path}")
        except Exception as e:
            self.logger.error(f"Error saving JSON: {str(e)}")
            raise

    def load_from_json(self, json_path: str) -> Dict[str, Any]:
        """从JSON文件加载数据"""
        try:
            with open(json_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            self.logger.error(f"Error loading JSON: {str(e)}")
            raise

    def merge_chord_runs(self, runs):
        merged_runs = merge_chord_runs(self, runs)
        return merged_runs 