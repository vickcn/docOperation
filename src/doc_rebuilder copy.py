from docx import Document
from docx.shared import Pt, Inches, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from typing import Dict, Any
import logging

class DocRebuilder:
    def __init__(self, logger=None):
        self.logger = logger or logging.getLogger(__name__)
        
    def rebuild_docx(self, data: Dict[str, Any], output_path: str) -> None:
        """
        从JSON数据重建docx文件
        """
        try:
            doc = Document()
            
            # 設置默認字體為 Calibri
            default_font = doc.styles['Normal'].font
            default_font.name = 'Calibri'
            
            # 重建樣式
            self._rebuild_styles(doc, data.get("styles", {}))
            
            # 重建元数据
            self._rebuild_metadata(doc, data.get("metadata", {}))
            
            # 重建段落
            self._rebuild_paragraphs(doc, data.get("paragraphs", []))
            
            # 重建表格
            self._rebuild_tables(doc, data.get("tables", []))
            
            # 保存文档
            doc.save(output_path)
            self.logger.info(f"Successfully rebuilt document: {output_path}")
        except Exception as e:
            self.logger.error(f"Error rebuilding document: {str(e)}")
            raise

    def _rebuild_styles(self, doc: Document, styles: Dict[str, Any]) -> None:
        """重建文档样式"""
        for style_name, style_data in styles.items():
            if style_name in doc.styles:
                style = doc.styles[style_name]
                if style_data.get("font_name"):
                    style.font.name = style_data["font_name"]
                if style_data.get("font_size"):
                    style.font.size = Pt(style_data["font_size"])
                if style_data.get("bold") is not None:
                    style.font.bold = style_data["bold"]
                if style_data.get("italic") is not None:
                    style.font.italic = style_data["italic"]

    def _rebuild_metadata(self, doc: Document, metadata: Dict[str, Any]) -> None:
        """重建文档元数据"""
        if "core_properties" in metadata:
            props = metadata["core_properties"]
            doc.core_properties.title = props.get("title", "")
            doc.core_properties.author = props.get("author", "")

    def _convert_to_safe_value(self, value: Any, default: int = 1) -> int:
        """将值转换为安全的整数范围"""
        if value is None:
            return default
        try:
            # 确保值在32位整数范围内
            safe_value = min(max(int(value), -2147483648), 2147483647)
            return safe_value
        except (ValueError, TypeError):
            return default

    def _apply_paragraph_format(self, paragraph, format_data: Dict[str, Any]) -> None:
        """应用段落格式"""
        if not format_data:
            return
            
        # 设置对齐方式
        if "alignment" in format_data:
            try:
                paragraph.alignment = format_data["alignment"]
            except ValueError:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        pf = paragraph.paragraph_format
        
        # 处理行距
        if "line_spacing" in format_data:
            try:
                line_spacing = format_data["line_spacing"]
                if isinstance(line_spacing, (int, float)):
                    # 如果是数字，转换为Twips
                    pf.line_spacing = Twips(int(line_spacing * 20))
                else:
                    # 默认使用单倍行距
                    pf.line_spacing = Twips(240)
            except Exception:
                pf.line_spacing = Twips(240)

        # 处理段落间距
        for spacing in ["space_before", "space_after"]:
            if spacing in format_data:
                try:
                    value = float(format_data[spacing])
                    setattr(pf, spacing, Pt(value))
                except (ValueError, TypeError):
                    pass

        # 处理缩进
        for indent in ["first_line_indent", "left_indent", "right_indent"]:
            if indent in format_data:
                try:
                    value = float(format_data[indent])
                    setattr(pf, indent, Pt(value))
                except (ValueError, TypeError):
                    pass

        # 处理段落控制
        for control in ["keep_together", "keep_with_next", "page_break_before", "widow_control"]:
            if control in format_data:
                setattr(pf, control, format_data[control])

    def _apply_run_format(self, run, format_data):
        """应用文本运行的格式"""
        if not format_data:
            return

        # 处理基本格式
        if format_data.get('bold') is not None:
            run.bold = format_data['bold']
            
        if format_data.get('italic') is not None:
            run.italic = format_data['italic']
            
        if format_data.get('underline') is not None:
            run.underline = format_data['underline']
            
        if format_data.get('font_size'):
            run.font.size = Pt(float(format_data['font_size']))
            
        # 处理字体设置
        rPr = run._element.rPr
        if rPr is None:
            rPr = run._element.get_or_add_rPr()
            
        # 确保 rFonts 元素存在
        rFonts = rPr.rFonts
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
            rPr.append(rFonts)
            
        # 应用各种字体设置
        if format_data.get('font_ascii'):
            rFonts.set(qn('w:ascii'), format_data['font_ascii'])
        if format_data.get('font_east_asia'):
            rFonts.set(qn('w:eastAsia'), format_data['font_east_asia'])
        if format_data.get('font_h_ansi'):
            rFonts.set(qn('w:hAnsi'), format_data['font_h_ansi'])
        if format_data.get('font_cs'):
            rFonts.set(qn('w:cs'), format_data['font_cs'])
        if format_data.get('font_hint'):
            rFonts.set(qn('w:hint'), format_data['font_hint'])

        # 处理上标和下标
        if format_data.get('subscript'):
            run.font.subscript = True
            
        if format_data.get('superscript'):
            run.font.superscript = True

        # 处理颜色
        if format_data.get('color'):
            color = format_data['color']
            if isinstance(color, str) and len(color) == 6:
                try:
                    r = int(color[0:2], 16)
                    g = int(color[2:4], 16)
                    b = int(color[4:6], 16)
                    run.font.color.rgb = RGBColor(r, g, b)
                except ValueError:
                    pass

        # 处理突出显示
        if format_data.get('highlight_color'):
            run._element.rPr.highlight = parse_xml(
                f'<w:highlight {nsdecls("w")} w:val="{format_data["highlight_color"]}"/>'
            )

        # 处理超链接
        if format_data.get('hyperlink'):
            self._apply_hyperlink(run, format_data['hyperlink'])

    def _rebuild_run(self, paragraph, run_data):
        """重建文本运行"""
        text = run_data.get('text', '')
        
        # 检查是否是和弦（包含升降号的情况）
        if text and text[0].isupper() and any(acc in text for acc in ['#', 'b']):
            # 分离和弦的基本音符和升降号
            base_note = text[0]  # 第一个字符是基本音符
            accidentals = text[1:]  # 剩余部分可能包含升降号和其他修饰符
            
            # 添加基本音符
            main_run = paragraph.add_run(base_note)
            self._apply_run_format(main_run, run_data)
            
            # 处理剩余部分
            i = 0
            while i < len(accidentals):
                if accidentals[i] in ['#', 'b']:
                    # 升降号只使用上标，保持原始格式
                    acc_run = paragraph.add_run(accidentals[i])
                    acc_format = run_data.copy()  # 复制原始格式
                    acc_format['superscript'] = True  # 只添加上标属性
                    self._apply_run_format(acc_run, acc_format)
                    i += 1
                else:
                    # 其他修饰符（如 m7）使用原始格式
                    modifier = accidentals[i:]
                    mod_run = paragraph.add_run(modifier)
                    self._apply_run_format(mod_run, run_data)
                    break
        else:
            # 非和弦文本，直接使用原始格式
            run = paragraph.add_run(text)
            self._apply_run_format(run, run_data)

    def rebuild_paragraph(self, doc, para_data):
        """重建段落"""
        # 创建新段落
        paragraph = doc.add_paragraph()
        
        # 应用段落样式
        if para_data.get('style'):
            paragraph.style = para_data['style']
            
        # 应用段落格式
        if para_data.get('format'):
            self._apply_paragraph_format(paragraph, para_data['format'])
            
        # 重建所有的文本运行
        for run_data in para_data.get('runs', []):
            self._rebuild_run(paragraph, run_data)
            
        return paragraph

    def _rebuild_paragraphs(self, doc: Document, paragraphs: list) -> None:
        """重建段落内容"""
        for para_data in paragraphs:
            paragraph = self.rebuild_paragraph(doc, para_data)

    def _rebuild_tables(self, doc: Document, tables: list) -> None:
        """重建表格内容"""
        for table_data in tables:
            if not table_data:
                continue
                
            rows = len(table_data)
            cols = len(table_data[0]) if rows > 0 else 0
            
            if rows > 0 and cols > 0:
                table = doc.add_table(rows=rows, cols=cols)
                for i, row_data in enumerate(table_data):
                    for j, cell_text in enumerate(row_data):
                        # 确保正确处理中文字符
                        text = str(cell_text)
                        if not isinstance(text, str):
                            text = str(text, 'utf-8', errors='replace')
                        table.cell(i, j).text = text 